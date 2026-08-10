"""Skeleton builder for `odoo-write-specifications` docx output.

Per-spec folder layout (created by the skill once Q3 brief lands and the
functionality slug is auto-derived):

    <repo>/specifications/<task-code> - <client> - <func>/
    ├── <task-code> - <client> - <func>.docx          ← the deliverable
    └── _reference/
        └── _build_<task-code>.py                      ← this file, customised

Run the per-task copy as:
    ./v19/odoo/.venv/bin/python \\
        "<repo>/specifications/<task-code> - <client> - <func>/_reference/_build_<task-code>.py"

The skill copies THIS file to the spec folder's `_reference/` subfolder per
task, then customises the SPEC_DATA block at the bottom with the interview
answers. The helpers in this module — make_cover, make_heading, make_table,
make_flow_strip, render_bpmn_image, add_phase_marker — stay shared across all
tasks.

Hand-editing the generated docx is NOT the workflow (P1). Edit SPEC_DATA, re-run.

Requires python-docx:
    ./v19/odoo/.venv/bin/pip install python-docx
"""
from __future__ import annotations

import argparse
import hashlib
import sys
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Cm, Inches, Mm, Pt, RGBColor
    from docx.enum.table import WD_ALIGN_VERTICAL
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_TAB_LEADER
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    sys.stderr.write(
        "python-docx not installed; run: ./v19/odoo/.venv/bin/pip install python-docx\n"
    )
    sys.exit(1)

# Pillow is OPTIONAL — only required when a workflow opts into the BPMN
# diagram via `bpmn_diagram` in SPEC_DATA. Workflows that use the flat
# flow_strip (or no process flow at all) need no extra dependency.
try:
    from PIL import Image, ImageDraw, ImageFont
    _PILLOW_AVAILABLE = True
except ImportError:
    _PILLOW_AVAILABLE = False


# ---------------------------------------------------------------------------
# Styling constants — mirrors docx_styling.md. Single source of truth.
#
# Aligned to Odoo's prescribed "Generic document" template: a monochrome grey
# heading ramp (#21272b → #434343 → #666666) set in Montserrat Medium, Open
# Sans body, and the Odoo brand plum (#714B67) as the single structural accent
# (table-header fills, rules, BPMN nodes). The only other colours are semantic:
# burnt-orange for Phase-2 / caution, green for success ends. See the Odoo
# template's word/styles.xml for the source values this mirrors.
# ---------------------------------------------------------------------------

# Odoo greyscale ink ramp (from the template's heading styles).
INK_STRONG = RGBColor(0x21, 0x27, 0x2B)   # Title, Heading 1–2, body
INK_MED = RGBColor(0x43, 0x43, 0x43)      # Heading 3
INK_MUTED = RGBColor(0x66, 0x66, 0x66)    # Heading 4, subtitle, captions, footers

# Odoo brand accent — the plum of the logo's first "o". Used for table-header
# fills, thin rules, first-column emphasis, and the BPMN shape language.
PLUM = RGBColor(0x71, 0x4B, 0x67)
PALE_PLUM = RGBColor(0xF1, 0xEC, 0xEF)    # ~6% plum tint — banner / alt-row fills

WHITE = RGBColor(0xFF, 0xFF, 0xFF)
BODY = INK_STRONG                          # near-black body text, Odoo ink
SECONDARY = INK_MUTED                      # captions / running header / footer
BORDER = RGBColor(0xB7, 0xB7, 0xB7)        # thin table border
PHASE2 = RGBColor(0xB4, 0x53, 0x09)        # burnt orange — Phase-2 / caution (semantic)
SUCCESS = RGBColor(0x15, 0x80, 0x3D)
WARN = PHASE2  # alias — used by the BPMN renderer for error end events

# Back-compat alias: older per-spec builder forks may reference NAVY. It now
# resolves to the Odoo plum accent so they pick up the new palette on re-run.
NAVY = PLUM
PALE_NAVY = PALE_PLUM

# RGB tuples for PIL drawing (PIL doesn't accept docx.shared.RGBColor directly).
# `accent` is the plum shape colour; `lane_label` is the solid sidebar band;
# `lane_alt` is the alternating swimlane background tint.
_BPMN_RGB = {
    "accent": (0x71, 0x4B, 0x67),
    "navy": (0x71, 0x4B, 0x67),            # back-compat key name; now plum
    "lane_label": (0x4A, 0x30, 0x43),      # darker plum — solid lane sidebar
    "lane_alt": (0xF7, 0xF3, 0xF5),        # faint plum tint — alternating lane fill
    "pale_navy": (0xF1, 0xEC, 0xEF),
    "ink": (0x21, 0x27, 0x2B),
    "body": (0x21, 0x27, 0x2B),
    "secondary": (0x66, 0x66, 0x66),
    "border": (0xC9, 0xC2, 0xC6),          # soft plum-grey gridline
    "lane_divider": (0xB3, 0xA4, 0xAE),    # lane separator line
    "warn": (0xB4, 0x53, 0x09),
    "success": (0x15, 0x80, 0x3D),
}

FONT_HEADING = "Montserrat Medium"         # Odoo template heading face
FONT_BODY = "Open Sans"
FONT_MONO = "JetBrains Mono"
FONT_BODY_FALLBACK = "Calibri"

# Type ramp — matches the Odoo template (Title 26 / H1 20 / H2 16 / H3 14 /
# H4 12, Open Sans 11 body). Tables run a touch tighter than body for density.
TITLE_SIZE = Pt(26)
SUBTITLE_SIZE = Pt(15)
HEADING_SIZES = {1: Pt(20), 2: Pt(16), 3: Pt(14), 4: Pt(12)}
HEADING_COLOURS = {1: INK_STRONG, 2: INK_STRONG, 3: INK_MED, 4: INK_MUTED}
BODY_SIZE = Pt(11)
TABLE_BODY_SIZE = Pt(10)
TABLE_HEADER_SIZE = Pt(10)
MONO_SIZE = Pt(9.5)

# The 11 per-workflow subsections, in canonical order — single source of truth
# for both the heading emitter (build) and the TOC (collect_toc), so the two
# can never drift. (SPEC_DATA key, heading label)
CANONICAL_SUBSECTIONS = [
    ("success_criteria", "Success Criteria"),
    ("user_stories_and_steps", "User Stories & Steps"),
    ("new_models", "New Models"),
    ("new_fields", "New Fields & Information"),
    ("navigation", "Navigation & Menus"),
    ("screens", "Screens & Interactions"),
    ("automated_behaviours", "Automated Behaviours"),
    ("business_rules", "Business Rules & Validations"),
    ("reports", "Reports & Analytics"),
    ("data_migration", "Data Migration"),
    ("access", "Access & Permissions"),
]


# ---------------------------------------------------------------------------
# Page / section setup
# ---------------------------------------------------------------------------

def _apply_font(run, name: str) -> None:
    """Bind a font family across all four OOXML script slots.

    python-docx's `run.font.name` only sets w:ascii + w:hAnsi. Heading faces
    like "Montserrat Medium" render more reliably (and match the Odoo template,
    which sets all four) when w:cs + w:eastAsia are bound too.
    """
    run.font.name = name
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    for slot in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rfonts.set(qn(slot), name)


def new_document() -> Document:
    doc = Document()
    section = doc.sections[0]
    section.page_height = Mm(297)
    section.page_width = Mm(210)
    # 25.4 mm (1 inch) margins on all four sides — matches Odoo's prescribed
    # document template (1440-twip margins). Header sits at the page edge,
    # footer 12.7 mm up. Usable column ~159 mm; fit-to-content tables keep wide
    # field tables inside it. See docx_styling.md § Page setup.
    section.top_margin = Mm(25.4)
    section.bottom_margin = Mm(25.4)
    section.left_margin = Mm(25.4)
    section.right_margin = Mm(25.4)
    section.header_distance = Mm(12.7)
    section.footer_distance = Mm(12.7)
    return doc


# ---------------------------------------------------------------------------
# Page furniture — Odoo logo header + confidentiality footer
#
# The Odoo "Generic document" template puts the logo in the running header and
# a thin-ruled centred footer. We adopt the logo + the rule, but the footer
# carries the spec's confidentiality + page numbering rather than a marketing
# URL (these are client functional specs). See docx_styling.md § Header / Footer.
# ---------------------------------------------------------------------------

def _find_logo(output_path: Path = None) -> Path | None:
    """Locate the embeddable Odoo logo PNG.

    The skill copies `reference/assets/odoo-logo.png` into each spec's
    `_reference/` folder at scaffold time, so the per-spec builder finds it
    as a sibling. When the skeleton runs in place we fall back to the skill's
    `reference/assets/`. Returns None if no logo is present — the build then
    proceeds without a logo rather than failing (graceful degradation, P6).
    """
    here = Path(__file__).resolve().parent
    candidates = [
        here / "odoo-logo.png",                       # per-spec _reference/ sibling
        here.parent / "assets" / "odoo-logo.png",     # skill reference/assets/
    ]
    if output_path is not None:
        candidates.insert(0, output_path.parent / "_reference" / "odoo-logo.png")
    for c in candidates:
        if c.is_file():
            return c
    return None


def _set_paragraph_rule(paragraph, *, top: bool = False, bottom: bool = False,
                        color_hex: str = "714B67", sz_eighths: int = 6) -> None:
    """Add a thin top and/or bottom border to a paragraph — the header/footer
    rule. sz is in eighths of a point per OOXML (6 = 0.75 pt)."""
    pPr = paragraph._p.get_or_add_pPr()
    pbdr = pPr.find(qn("w:pBdr"))
    if pbdr is None:
        pbdr = OxmlElement("w:pBdr")
        pPr.append(pbdr)
    for edge, want in (("top", top), ("bottom", bottom)):
        if not want:
            continue
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), str(sz_eighths))
        el.set(qn("w:space"), "4")
        el.set(qn("w:color"), color_hex)
        pbdr.append(el)


def _add_logo(paragraph, logo_path: Path, height_mm: float) -> None:
    """Add the Odoo logo as an inline image in `paragraph`, preserving aspect."""
    run = paragraph.add_run()
    run.add_picture(str(logo_path), height=Mm(height_mm))


def setup_page_furniture(doc: Document, spec_data: dict,
                         output_path: Path = None) -> None:
    """Wire the running header (task-code text + Odoo logo, thin rule below) and
    the footer (confidentiality + Page N of T, thin rule above) onto the section.

    Uses a different first page so the cover carries no running header (its logo
    sits in the cover body) while still showing the footer."""
    section = doc.sections[0]
    section.different_first_page_header_footer = True

    task_code = spec_data.get("task_code", "")
    client = spec_data.get("client_name", "")
    md = spec_data.get("metadata", {})
    consultancy = md.get("consultancy") or _consultancy_from_role(md.get("role", ""))
    content_w_mm = 210 - 25.4 * 2  # usable width inside 1-inch margins
    logo = _find_logo(output_path)

    # --- Default header (page 2+): running text left, logo right, rule below ---
    hdr = section.header
    hdr.is_linked_to_previous = False
    hp = hdr.paragraphs[0]
    hp.text = ""
    hp.paragraph_format.tab_stops.add_tab_stop(Mm(content_w_mm), WD_TAB_ALIGNMENT.RIGHT)
    label = " — ".join([x for x in (task_code, client, "Functional Specification") if x])
    run = hp.add_run(label)
    run.italic = True
    run.font.name = FONT_BODY
    run.font.size = Pt(8)
    run.font.color.rgb = INK_MUTED
    if logo:
        hp.add_run("\t")
        _add_logo(hp, logo, height_mm=6.5)
    _set_paragraph_rule(hp, bottom=True)

    # First-page header stays empty (cover already shows the logo prominently).
    fp_hdr = section.first_page_header
    fp_hdr.is_linked_to_previous = False
    fp_hdr.paragraphs[0].text = ""

    # --- Footer (both pages): confidentiality + Page N of T, rule above ---
    conf = f"Confidential — {client}" if client else "Confidential"
    if consultancy:
        conf += f" & {consultancy}"

    def _build_footer(footer):
        footer.is_linked_to_previous = False
        fpar = footer.paragraphs[0]
        fpar.text = ""
        fpar.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_paragraph_rule(fpar, top=True)
        r = fpar.add_run(conf + "    |    Page ")
        r.font.name = FONT_BODY
        r.font.size = Pt(8.5)
        r.font.color.rgb = INK_MUTED
        _insert_field(fpar, "PAGE", font_size_pt=8.5, italic=False, color_rgb=INK_MUTED)
        r = fpar.add_run(" of ")
        r.font.name = FONT_BODY
        r.font.size = Pt(8.5)
        r.font.color.rgb = INK_MUTED
        _insert_field(fpar, "NUMPAGES", font_size_pt=8.5, italic=False, color_rgb=INK_MUTED)

    _build_footer(section.footer)
    _build_footer(section.first_page_footer)


def _consultancy_from_role(role: str) -> str:
    """Best-effort consultancy name from a 'Title, Firm' role string."""
    if role and "," in role:
        return role.split(",")[-1].strip()
    return ""


# ---------------------------------------------------------------------------
# Cover page
# ---------------------------------------------------------------------------

def make_cover(doc: Document, title: str, subtitle: str, metadata: dict,
               logo_path: Path = None) -> None:
    """Render the cover page per docx_styling.md.

    Odoo template treatment: centred Odoo logo at top, then the title in
    Montserrat Medium 26 pt (#21272b), a grey subtitle (15 pt), a thin plum
    rule, then the left-indented metadata block.

    metadata keys (all required except `companion_to`):
        prepared_by, role, client, companion_to (optional), scope, date, version
    """
    for _ in range(4):
        doc.add_paragraph()

    if logo_path is not None:
        lp = doc.add_paragraph()
        lp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        lp.paragraph_format.space_after = Pt(18)
        _add_logo(lp, logo_path, height_mm=22)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(title)
    run.font.size = TITLE_SIZE
    run.font.color.rgb = INK_STRONG
    _apply_font(run, FONT_HEADING)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(subtitle)
    run.font.size = SUBTITLE_SIZE
    run.font.color.rgb = INK_MUTED
    _apply_font(run, FONT_BODY)

    # Thin plum rule under the subtitle (drawn as a bottom-bordered empty
    # paragraph — renders reliably across Word / Pages / LibreOffice).
    rule = doc.add_paragraph()
    rule.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rule.paragraph_format.space_before = Pt(6)
    rule.paragraph_format.space_after = Pt(18)
    _set_paragraph_rule(rule, bottom=True, color_hex="714B67", sz_eighths=8)

    meta_lines = [
        ("Prepared by", metadata["prepared_by"]),
        ("Role", metadata["role"]),
        ("Client", metadata["client"]),
    ]
    if metadata.get("companion_to"):
        meta_lines.append(("Companion to", metadata["companion_to"]))
    meta_lines.extend([
        ("Scope", metadata["scope"]),
        ("Date", metadata["date"]),
        ("Version", metadata["version"]),
    ])

    for label, value in meta_lines:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.left_indent = Cm(5)
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(f"{label}:  ")
        run.bold = True
        run.font.size = BODY_SIZE
        run.font.name = FONT_BODY
        run.font.color.rgb = INK_STRONG
        run = p.add_run(value)
        run.font.size = BODY_SIZE
        run.font.name = FONT_BODY
        run.font.color.rgb = INK_STRONG

    doc.add_page_break()


# ---------------------------------------------------------------------------
# Headings + body
# ---------------------------------------------------------------------------

def make_heading(doc: Document, text: str, level: int = 1) -> None:
    """Render a styled heading.

    Binds the Word built-in `Heading {level}` paragraph style so the Stage-1
    linter can identify sections by style and Word's outline navigation works.
    Levels 1–3 are also bookmarked (see _add_bookmark) so the populated TOC can
    hyperlink to them and resolve their page numbers. Run-level formatting
    (Montserrat Medium, the Odoo grey ink ramp, template sizes) overrides the
    style's defaults while keeping the style binding.

    Per the Odoo template the heading face is Montserrat Medium and the weight
    comes from the "Medium" face itself, so `run.bold` stays off — matching the
    template's heading rPr exactly. Colour ramps by level: #21272b (1–2),
    #434343 (3), #666666 (4).
    """
    p = doc.add_paragraph()
    try:
        p.style = doc.styles[f"Heading {level}"]
    except KeyError:
        pass
    if level == 1:
        p.paragraph_format.space_before = Pt(20)
        p.paragraph_format.space_after = Pt(6)
    elif level == 2:
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(6)
    else:
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.size = HEADING_SIZES.get(level, Pt(11))
    run.font.color.rgb = HEADING_COLOURS.get(level, INK_STRONG)
    _apply_font(run, FONT_HEADING)
    # Bookmark levels 1–3 so the populated TOC (make_toc) can link to them and
    # resolve their page numbers (see _toc_bm_name / make_toc).
    if level <= 3:
        _add_bookmark(p, _toc_bm_name(text))


def make_paragraph(doc: Document, text: str, bold: bool = False, mono: bool = False) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(text)
    run.bold = bold
    run.font.size = MONO_SIZE if mono else BODY_SIZE
    run.font.name = FONT_MONO if mono else FONT_BODY
    run.font.color.rgb = BODY


def make_bullet(doc: Document, text: str) -> None:
    """Bulleted paragraph with a literal U+2022 (•) bullet character.

    Why not `doc.add_paragraph(style="List Bullet")`: that style references
    the Symbol font for the bullet glyph, and on macOS Pages (plus some
    Word versions) the Symbol fallback fails — the bullet renders as an
    empty rectangle. The U+2022 codepoint is supported by every modern
    body font (Open Sans, Calibri, JetBrains Mono), so it renders
    consistently across Word, Pages, and LibreOffice. Indentation is set
    manually via paragraph_format so the visual result matches the List
    Bullet style without relying on its font reference.
    """
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.left_indent = Mm(6)
    p.paragraph_format.first_line_indent = Mm(-6)
    run = p.add_run("•\t")
    run.font.name = FONT_BODY
    run.font.size = BODY_SIZE
    run.font.color.rgb = BODY
    run = p.add_run(text)
    run.font.name = FONT_BODY
    run.font.size = BODY_SIZE
    run.font.color.rgb = BODY


# ---------------------------------------------------------------------------
# Tables
# ---------------------------------------------------------------------------

def _set_cell_fill(cell, hex_color: str) -> None:
    """Set a table-cell background fill."""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _set_cell_border(cell, color_hex: str = "B7B7B7", sz_pt: int = 4) -> None:
    """Thin grey border on all four sides."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        border = OxmlElement(f"w:{side}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), str(sz_pt))
        border.set(qn("w:color"), color_hex)
        tcBorders.append(border)
    tcPr.append(tcBorders)


def _set_cell_margins(cell, top_pt: int = 2, left_pt: int = 4) -> None:
    """Tight cell padding so tables fit content. 2 pt vertical, 4 pt horizontal."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement("w:tcMar")
    # Word stores margins in twentieths of a point (1 pt = 20 twips).
    for side, pt in (("top", top_pt), ("bottom", top_pt),
                    ("left", left_pt), ("right", left_pt)):
        m = OxmlElement(f"w:{side}")
        m.set(qn("w:w"), str(pt * 20))
        m.set(qn("w:type"), "dxa")
        tcMar.append(m)
    tcPr.append(tcMar)


def _set_table_fit_content(table) -> None:
    """Make the table truly AutoFit-to-Contents.

    python-docx's `table.autofit = True` writes ambiguous XML; Word treats
    the resulting table as AutoFit-to-Window (equal-width columns spanning
    the page). Setting only `tblLayout=autofit` + `tblW w=0 type=auto` is
    not enough either — python-docx ALSO writes equal `<w:gridCol w:w="N"/>`
    elements in `<w:tblGrid>` and explicit `<w:tcW w:w="N" w:type="dxa"/>`
    on every cell. Those starting widths override the autofit hint; Word
    treats them as the desired widths and only nudges them slightly.

    True AutoFit-to-Contents needs all three pieces:

        <w:tblPr>
          <w:tblW w:w="0" w:type="auto"/>          ← preferred width = auto
          <w:tblLayout w:type="autofit"/>           ← autofit algorithm
        </w:tblPr>
        <w:tblGrid>
          <w:gridCol w:w="0"/>                      ← no preferred col width
          ...
        </w:tblGrid>
        <!-- per cell: -->
        <w:tcW w:w="0" w:type="auto"/>              ← no preferred cell width

    With all three, Word measures each column's longest content line and
    sizes the column accordingly: short-value columns ("Sales", "1",
    "Sales rep") collapse to their natural narrow width; prose columns
    absorb the remainder.
    """
    tbl = table._tbl
    tblPr = tbl.tblPr
    # 1. tblPr — table fills 100% of page text width; cell widths auto-
    # distribute by content. NOT tblW=0/auto: that lets the table collapse
    # to its minimum content width, producing tables that occupy ~25% of
    # the page with huge whitespace to the right and text wrapping per
    # syllable inside narrow columns. We want "AutoFit to Window with
    # content-proportional columns" — i.e. table=100% page, columns=auto.
    for tag in ("w:tblW", "w:tblLayout"):
        for child in tblPr.findall(qn(tag)):
            tblPr.remove(child)
    tblW = OxmlElement("w:tblW")
    tblW.set(qn("w:w"), "5000")   # 5000 fiftieths-of-a-percent = 100%
    tblW.set(qn("w:type"), "pct")
    tblPr.append(tblW)
    tblLayout = OxmlElement("w:tblLayout")
    tblLayout.set(qn("w:type"), "autofit")
    tblPr.append(tblLayout)
    # 2. tblGrid — zero out every gridCol's preferred width
    tblGrid = tbl.find(qn("w:tblGrid"))
    if tblGrid is not None:
        for gridCol in tblGrid.findall(qn("w:gridCol")):
            gridCol.set(qn("w:w"), "0")
    # 3. Per-cell tcW — set to 0/auto so each cell defers to autofit
    for tr in tbl.findall(qn("w:tr")):
        for tc in tr.findall(qn("w:tc")):
            tcPr = tc.find(qn("w:tcPr"))
            if tcPr is None:
                continue
            tcW = tcPr.find(qn("w:tcW"))
            if tcW is None:
                tcW = OxmlElement("w:tcW")
                tcPr.insert(0, tcW)  # tcW must come first per OOXML schema
            tcW.set(qn("w:w"), "0")
            tcW.set(qn("w:type"), "auto")


def _insert_field(paragraph, instr_text: str, *, placeholder: str = "1",
                  font_size_pt: int = 8, italic: bool = True,
                  color_rgb: RGBColor = None) -> None:
    """Insert a Word field (PAGE / NUMPAGES / TOC) into a paragraph correctly.

    A working Word field needs four fldChar/instrText children in order:
    begin → instrText → **separate** → (placeholder text) → end. The
    `separate` fldChar is the one most hand-rolled implementations miss;
    without it, Word and Pages render the field as blank or as the raw
    instrText. The placeholder text sits between separate and end and
    holds whatever Word/Pages overwrites on render.

    `instr_text` is the field code BODY (no surrounding spaces). Common
    values: "PAGE", "NUMPAGES", 'TOC \\o "1-3" \\h \\z \\u'.
    """
    if color_rgb is None:
        color_rgb = SECONDARY

    def _styled_run():
        r = paragraph.add_run()
        r.font.name = FONT_BODY
        r.font.size = Pt(font_size_pt)
        r.italic = italic
        r.font.color.rgb = color_rgb
        return r

    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = f" {instr_text} "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")

    r1 = _styled_run(); r1._r.append(begin)
    r2 = _styled_run(); r2._r.append(instr)
    r3 = _styled_run(); r3._r.append(separate)
    # Placeholder run — Word/Pages overwrites this with the computed value.
    placeholder_run = _styled_run()
    placeholder_run.add_text(placeholder)
    r5 = _styled_run(); r5._r.append(end)


def _insert_page_field(paragraph, field_name: str, *, font_size_pt: int = 8,
                       italic: bool = True, color_rgb: RGBColor = None) -> None:
    """Back-compat wrapper around _insert_field for PAGE / NUMPAGES."""
    _insert_field(
        paragraph, field_name, placeholder="1",
        font_size_pt=font_size_pt, italic=italic, color_rgb=color_rgb,
    )


_BM_ID = [1000]  # unique numeric ids for heading bookmarks


def _toc_bm_name(text: str) -> str:
    """Stable, valid, unique Word bookmark name derived from heading text.
    The same text → same name, so make_heading's bookmark and make_toc's
    PAGEREF/hyperlink always point at each other."""
    return "_Toc" + hashlib.md5(text.encode("utf-8")).hexdigest()[:12]


def _add_bookmark(paragraph, name: str) -> None:
    """Wrap a heading paragraph's content in a Word bookmark so a TOC
    PAGEREF / hyperlink can target it."""
    _BM_ID[0] += 1
    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), str(_BM_ID[0]))
    start.set(qn("w:name"), name)
    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), str(_BM_ID[0]))
    paragraph._p.insert(0, start)
    paragraph._p.append(end)


def _toc_rpr(size_pt: int, bold: bool, color_hex: str):
    """A styled w:rPr element for a TOC run (title or page number)."""
    rpr = OxmlElement("w:rPr")
    rfonts = OxmlElement("w:rFonts")
    rfonts.set(qn("w:ascii"), FONT_BODY)
    rfonts.set(qn("w:hAnsi"), FONT_BODY)
    rpr.append(rfonts)
    if bold:
        rpr.append(OxmlElement("w:b"))
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), str(int(size_pt * 2)))  # half-points
    rpr.append(sz)
    col = OxmlElement("w:color")
    col.set(qn("w:val"), color_hex)
    rpr.append(col)
    return rpr


def _pageref_run_el(bm_name: str, size_pt: int):
    """A fully-built w:r containing a PAGEREF field (cached '1', refreshed on
    open via the document's updateFields flag) pointing at a heading bookmark."""
    r = OxmlElement("w:r")
    r.append(_toc_rpr(size_pt, False, "1A1A1A"))
    for kind in ("begin", "instr", "separate", "text", "end"):
        if kind in ("begin", "separate", "end"):
            fld = OxmlElement("w:fldChar")
            fld.set(qn("w:fldCharType"), kind)
            r.append(fld)
        elif kind == "instr":
            instr = OxmlElement("w:instrText")
            instr.set(qn("xml:space"), "preserve")
            instr.text = f" PAGEREF {bm_name} \\h "
            r.append(instr)
        else:
            t = OxmlElement("w:t")
            t.text = "1"
            r.append(t)
    return r


def collect_toc(spec_data) -> list:
    """Ordered (level, heading_text) entries for TOC levels 1–3, derived from
    the SAME structure build() emits (and gated by the same conditionals), so
    the populated TOC can never drift from the headings. Pass the result to
    make_toc."""
    entries = [
        (1, "1. Odoo Version"),
        (1, "2. Business Case"),
        (1, "3. Apps Impacted / New Apps Proposed"),
    ]
    if spec_data.get("apps_impacted"):
        entries.append((2, "3.1 Impacted Standard Apps"))
    if spec_data.get("new_apps"):
        entries.append((2, "3.2 New Custom Apps"))
    entries.append((1, "4. Functional Layout — Workflows"))
    entries.append((1, "5. Per-Workflow Detail"))
    for w_idx, w in enumerate(spec_data["workflows"], start=1):
        entries.append((2, f"5.{w_idx} {w['name']}"))
        for sub_idx, (_key, heading) in enumerate(CANONICAL_SUBSECTIONS, start=1):
            entries.append((3, f"5.{w_idx}.{sub_idx} {heading}"))
    return entries


def set_update_fields(doc: Document) -> None:
    """Tell Word/Pages/LibreOffice to refresh all fields (TOC page numbers) on
    open, so the TOC shows correct numbers with no manual 'Update Field'."""
    settings = doc.settings.element
    if settings.find(qn("w:updateFields")) is None:
        upd = OxmlElement("w:updateFields")
        upd.set(qn("w:val"), "true")
        settings.append(upd)


def make_toc(doc: Document, entries=None, *, title: str = "Table of Contents",
             max_level: int = 3) -> None:
    """Render a POPULATED Table of Contents — titles + page numbers visible on
    open, never a 'right-click to update' placeholder.

    Pass `entries` as an ordered list of `(level, heading_text)` tuples (levels
    1–`max_level`), where each `heading_text` EXACTLY matches the string passed
    to `make_heading` (so bookmark names line up). Each entry renders as a
    hyperlink to the heading's bookmark + a dot-leader tab + a PAGEREF page
    number. Titles are real text (visible in any viewer); page numbers refresh
    on open because the document carries `w:updateFields` (see set_update_fields,
    which this calls).

    The per-task builder derives `entries` from its SPEC_DATA structure (the
    same ordering build() emits) — keep a single source of truth for the
    subsection labels so the TOC can't drift from the headings.

    Fallback: if `entries` is None, insert a live Word TOC field (still with
    updateFields set, so Word/Pages/LibreOffice populate it on open). Prefer
    passing `entries` so the TOC is populated even in viewers that don't refresh
    fields.
    """
    set_update_fields(doc)
    make_heading(doc, title, level=1)

    if entries is None:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.15
        _insert_field(
            p, f'TOC \\o "1-{max_level}" \\h \\z \\u',
            placeholder="Updating Table of Contents…",
            font_size_pt=10, italic=False, color_rgb=BODY,
        )
        doc.add_page_break()
        return

    indent = {1: 0, 2: 14, 3: 28}
    ink_hex = "21272B"  # Odoo strong ink for level-1 TOC entries
    for level, text in entries:
        bm = _toc_bm_name(text)
        size = 10 if level == 1 else 9
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.left_indent = Pt(indent.get(level, 0))
        # Right-aligned dot-leader tab at the content-width edge (159 mm inside
        # the 1-inch A4 margins).
        p.paragraph_format.tab_stops.add_tab_stop(
            Mm(159), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)

        hyperlink = OxmlElement("w:hyperlink")
        hyperlink.set(qn("w:anchor"), bm)
        title_run = OxmlElement("w:r")
        title_run.append(_toc_rpr(size, level == 1,
                                  ink_hex if level == 1 else "434343"))
        t = OxmlElement("w:t")
        t.set(qn("xml:space"), "preserve")
        t.text = text
        title_run.append(t)
        tab_run = OxmlElement("w:r")
        tab_run.append(_toc_rpr(size, False, "1A1A1A"))
        tab_run.append(OxmlElement("w:tab"))
        hyperlink.append(title_run)
        hyperlink.append(tab_run)
        p._p.append(hyperlink)
        p._p.append(_pageref_run_el(bm, size))

    doc.add_page_break()


def make_table(
    doc: Document,
    rows: list[list[str]],
    header_row: bool = True,
    first_col_emphasis: bool = True,
) -> None:
    if not rows:
        return
    cols = len(rows[0])
    table = doc.add_table(rows=len(rows), cols=cols)
    table.alignment = WD_ALIGN_PARAGRAPH.LEFT
    # True AutoFit-to-Contents via XML — python-docx's `table.autofit`
    # setter is unreliable; see _set_table_fit_content docstring.
    _set_table_fit_content(table)

    for r_idx, row in enumerate(rows):
        for c_idx, text in enumerate(row):
            cell = table.cell(r_idx, c_idx)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
            cell.text = ""
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.line_spacing = 1.15
            run = p.add_run(text)
            run.font.name = FONT_BODY
            run.font.size = TABLE_HEADER_SIZE if (header_row and r_idx == 0) else TABLE_BODY_SIZE

            if header_row and r_idx == 0:
                _set_cell_fill(cell, "714B67")  # Odoo plum accent
                run.bold = True
                run.font.color.rgb = WHITE
            else:
                run.font.color.rgb = BODY
                if first_col_emphasis and c_idx == 0:
                    run.bold = True
                    run.font.color.rgb = PLUM

            _set_cell_border(cell)
            _set_cell_margins(cell)


# ---------------------------------------------------------------------------
# Flow strip — per-workflow opener
# ---------------------------------------------------------------------------

def make_flow_strip(doc: Document, steps: list[str]) -> None:
    """Legacy table-based flow strip — refreshed pill-and-arrow look.

    Used as the fallback path when Pillow is unavailable in the venv
    (bare-Python install, user declined the Pillow install gate). The
    primary path is `render_flow_strip_image` + `make_flow_strip_diagram`
    — a PIL-rendered PNG with rounded pills, numbered step badges, and
    clean arrows, visually consistent with the BPMN diagram language.

    This table version mirrors the PIL renderer's visual vocabulary
    within docx table constraints:
      - White cell fill + navy outline + generous padding → "pill" feel
        (docx tables can't do true rounded corners, but the white-on-
        navy-border read is close)
      - Numbered prefix on each step ("1. ", "2. ", …) — matches the
        PIL renderer's filled-circle step-badge convention
      - `→` glyph (U+2192) instead of `▶` — reads as a flow connector,
        not a play button. Rendered larger (16pt) on a transparent /
        unbordered cell so it floats between the pills.
    """
    if not steps:
        return
    cells_count = len(steps) * 2 - 1
    table = doc.add_table(rows=1, cols=cells_count)
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Plum hex for the cell border (the docx-table substitute for the
    # PIL renderer's plum outline on the rounded pill).
    NAVY_HEX = "714B67"

    for i, step in enumerate(steps):
        col_idx = i * 2
        cell = table.cell(0, col_idx)
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # Numbered prefix — same step-badge convention as the PIL renderer.
        badge_run = p.add_run(f"{i + 1}. ")
        badge_run.bold = True
        badge_run.font.name = FONT_BODY
        badge_run.font.size = Pt(11)
        badge_run.font.color.rgb = NAVY
        # Step label
        label_run = p.add_run(step)
        label_run.bold = True
        label_run.font.name = FONT_BODY
        label_run.font.size = Pt(11)
        label_run.font.color.rgb = NAVY
        # White fill + navy 1pt border = the pill look within docx constraints.
        # (sz_pt is in eighths of a point per OOXML, so 8 = 1pt.)
        _set_cell_fill(cell, "FFFFFF")
        _set_cell_border(cell, color_hex=NAVY_HEX, sz_pt=8)
        _set_cell_margins(cell, top_pt=8, left_pt=12)

        if i < len(steps) - 1:
            arrow_cell = table.cell(0, col_idx + 1)
            arrow_cell.text = ""
            p = arrow_cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run("→")
            run.font.name = FONT_BODY
            run.font.size = Pt(16)
            run.font.color.rgb = NAVY
            # Arrow cell stays transparent (no fill, no border) so the
            # → reads as floating between adjacent pills rather than
            # boxed in. No _set_cell_fill / _set_cell_border call here.


def render_flow_strip_image(steps, output_path: Path, title: str = None) -> Path:
    """Render a modern block-diagram (flow strip) as a PNG.

    Each step renders as a rounded-rectangle pill with a numbered step
    badge in the top-left, the step label centred (with word-wrap), and
    a clean arrow connecting it to the next pill. Visually consistent
    with the BPMN diagram language: same navy palette, same step-badge
    convention, same shape vocabulary.

    `steps` accepts EITHER:
      - a list of plain strings (each becomes a pill labelled with the
        string — the standard SPEC_DATA shape) OR
      - a list of dicts `{"label": str, "description"?: str}` for
        richer per-step content (future use; description currently
        ignored — the pill height isn't sized for two-line content yet).

    `title` is optional. When provided, renders above the pill row in
    navy 16pt. Workflow-name headings already sit above the diagram in
    the docx, so the title is usually redundant — leave `None` unless
    the diagram is being used standalone.

    Returns `output_path` (matches the BPMN renderer's signature for
    chaining: `make_flow_strip_diagram(doc, render_flow_strip_image(...))`).
    """
    if not _PILLOW_AVAILABLE:
        raise ImportError(
            "Pillow is required for the modern flow strip renderer; "
            "install with `./v19/odoo/.venv/bin/pip install Pillow` or use "
            "the legacy table-based `make_flow_strip()` fallback."
        )

    # Normalize input
    normalized = []
    for s in steps:
        if isinstance(s, str):
            normalized.append({"label": s})
        elif isinstance(s, dict) and "label" in s:
            normalized.append(s)
        else:
            raise ValueError(
                f"flow_strip entry must be a string or a dict with a "
                f"'label' key; got {s!r}"
            )
    n = len(normalized)
    if n == 0:
        raise ValueError("flow_strip requires at least one step")

    # Layout constants (px @ _BPMN_DPI) — pill geometry shares the BPMN
    # renderer's visual vocabulary (rounded corners, numbered badges, plum
    # palette) and the same readable-on-page sizing.
    PILL_W = 332
    PILL_H = 138
    PILL_RADIUS = 18
    ARROW_GAP = 92        # horizontal space between adjacent pills
    LEFT_MARGIN = 44
    TOP_MARGIN = 44
    BOTTOM_MARGIN = 44
    TITLE_BLOCK_H = 56 if title else 0

    W = LEFT_MARGIN + n * PILL_W + (n - 1) * ARROW_GAP + LEFT_MARGIN
    H = TOP_MARGIN + TITLE_BLOCK_H + PILL_H + BOTTOM_MARGIN

    img = Image.new("RGB", (W, H), "white")
    draw = ImageDraw.Draw(img)
    R = _BPMN_RGB

    # Fonts (px)
    title_font = _resolve_bpmn_font(34) if title else None
    label_font = _resolve_bpmn_font(30)
    badge_font = _resolve_bpmn_font(26)

    # Optional title
    if title:
        draw.text((LEFT_MARGIN, 18), title, fill=R["accent"], font=title_font)

    pill_y_top = TOP_MARGIN + TITLE_BLOCK_H
    pill_y_bottom = pill_y_top + PILL_H
    pill_cy = pill_y_top + PILL_H // 2

    def _draw_pill_label(cx, cy, text, font, colour, max_width):
        """Wrap and centre the step label within the pill body."""
        words = text.split()
        if not words:
            return
        lines = []
        current = [words[0]]
        for w in words[1:]:
            trial = " ".join(current + [w])
            bbox = draw.textbbox((0, 0), trial, font=font)
            if (bbox[2] - bbox[0]) <= max_width:
                current.append(w)
            else:
                lines.append(" ".join(current))
                current = [w]
        lines.append(" ".join(current))

        ascent = draw.textbbox((0, 0), "Ag", font=font)
        line_h = (ascent[3] - ascent[1]) + 8
        total_h = line_h * len(lines)
        y = cy - total_h // 2 + 4
        for line in lines:
            bbox = draw.textbbox((0, 0), line, font=font)
            tw = bbox[2] - bbox[0]
            draw.text((cx - tw // 2, y), line, fill=colour, font=font)
            y += line_h

    def _draw_flow_arrow(x1, y, x2):
        """Horizontal arrow with a filled-triangle arrowhead."""
        draw.line([(x1, y), (x2, y)], fill=R["accent"], width=3)
        head = 16
        draw.polygon([
            (x2, y),
            (x2 - head, y - 8),
            (x2 - head, y + 8),
        ], fill=R["accent"])

    # Pills + arrows
    for i, step in enumerate(normalized):
        pill_x_left = LEFT_MARGIN + i * (PILL_W + ARROW_GAP)
        pill_x_right = pill_x_left + PILL_W
        pill_cx = pill_x_left + PILL_W // 2

        # Pill body (rounded rectangle, plum outline, white fill) with an
        # accent top-stripe — the same recognisably-Odoo card feel as a task.
        draw.rounded_rectangle(
            [(pill_x_left, pill_y_top), (pill_x_right, pill_y_bottom)],
            radius=PILL_RADIUS,
            fill="white",
            outline=R["accent"],
            width=3,
        )
        draw.rounded_rectangle(
            [(pill_x_left, pill_y_top), (pill_x_right, pill_y_top + 14)],
            radius=7, fill=R["accent"], outline=R["accent"],
        )

        # Numbered step badge (top-left, filled plum circle, white number)
        badge_r = 19
        badge_cx = pill_x_left + badge_r + 14
        badge_cy = pill_y_top + badge_r + 18
        draw.ellipse(
            [(badge_cx - badge_r, badge_cy - badge_r),
             (badge_cx + badge_r, badge_cy + badge_r)],
            fill=R["accent"], outline="white", width=2,
        )
        bbox = draw.textbbox((0, 0), str(i + 1), font=badge_font)
        tw = bbox[2] - bbox[0]
        th = bbox[3] - bbox[1]
        draw.text(
            (badge_cx - tw // 2, badge_cy - th // 2 - bbox[1]),
            str(i + 1), fill="white", font=badge_font,
        )

        # Step label — centred, wrapped to fit pill width minus padding
        _draw_pill_label(
            pill_cx, pill_cy + 8, step["label"],
            label_font, R["ink"],
            max_width=PILL_W - 44,
        )

        # Arrow to next pill
        if i < n - 1:
            arrow_x1 = pill_x_right + 10
            arrow_x2 = pill_x_right + ARROW_GAP - 10
            _draw_flow_arrow(arrow_x1, pill_cy, arrow_x2)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    img.save(output_path, format="PNG", dpi=(_BPMN_DPI, _BPMN_DPI))
    return output_path


def make_flow_strip_diagram(doc: Document, image_path: Path,
                            caption: str = None) -> None:
    """Embed a flow strip PNG into the docx, centred, full content width.

    Companion to `render_flow_strip_image`. Same embed pattern as
    `make_bpmn_diagram`. Caller ensures `image_path` exists (typically
    via `render_flow_strip_image` immediately prior).
    """
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run()
    run.add_picture(str(image_path), width=_diagram_width(image_path))
    if caption:
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.paragraph_format.space_after = Pt(8)
        cap_run = cap.add_run(caption)
        cap_run.italic = True
        cap_run.font.name = FONT_BODY
        cap_run.font.size = Pt(9)
        cap_run.font.color.rgb = SECONDARY


# ---------------------------------------------------------------------------
# BPMN swimlane diagram (data-driven, optional richer alternative to
# make_flow_strip). Renders a PNG via Pillow and embeds it via
# `make_bpmn_diagram`. Workflows opt in by populating a `bpmn_diagram`
# block in their workflow dict (see § BPMN diagram spec in SKILL.md for
# the schema). Workflows that don't include the block fall back to the
# flat `flow_strip` rendering.
# ---------------------------------------------------------------------------

# The BPMN/flow canvas is treated at this PPI when embedded in the docx, so a
# 31-px label lands near body scale on the page (see make_bpmn_diagram). Render
# at this density rather than the old "1500-px canvas squeezed to 6.5 in" which
# crushed labels to ~3.5 pt — the "font almost unreadable" complaint.
_BPMN_DPI = 220
_BPMN_CONTENT_W_IN = 6.27       # usable width inside 1-inch A4 margins
_FONT_CACHE: dict = {}
# Prefer the document's own faces if the host has them, else clean sans
# fallbacks. Open Sans / Montserrat keep the diagram type consistent with the
# docx body; Helvetica / Arial / DejaVu are the cross-platform safety net.
_FONT_CANDIDATES = [
    "/Library/Fonts/OpenSans-Regular.ttf",
    "/Users/Shared/Fonts/OpenSans-Regular.ttf",
    "/Library/Fonts/Montserrat-Medium.ttf",
    "/System/Library/Fonts/Helvetica.ttc",
    "/System/Library/Fonts/Supplemental/Arial.ttf",
    "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
]


def _resolve_bpmn_font(size_px: int):
    """Locate a usable TrueType font for diagram labels at the given pixel size.

    Cached per size. Falls back through Open Sans / Montserrat (matching the
    docx body) then the OS sans defaults; final fallback is Pillow's bitmap
    default so the diagram still renders on a host with no TrueType fonts.
    """
    key = int(size_px)
    if key in _FONT_CACHE:
        return _FONT_CACHE[key]
    for path in _FONT_CANDIDATES:
        try:
            f = ImageFont.truetype(path, key)
            _FONT_CACHE[key] = f
            return f
        except (OSError, IOError):
            continue
    f = ImageFont.load_default()
    _FONT_CACHE[key] = f
    return f


def render_bpmn_image(diagram_data: dict, output_path: Path,
                      draw_header: bool = True) -> Path:
    """Render a BPMN-style swimlane diagram as a PNG.

    `diagram_data` schema:

        title (str)         — large title above the diagram
        subtitle (list[str]) — 1–3 lines of context, optional
        lanes (list[dict])   — [{"name": str, "sub": str}, ...]
        nodes (list[dict])   — see node schema below
        edges (list[dict])   — [{"from": id, "to": id,
                                 "label"?: str, "label_colour"?: "success"|"warn"}, ...]

    Node schema (every node has "id", "kind", "lane", "col"):
        kind = "start_event"  → "label" (str, rendered below the circle)
        kind = "end_event"    → "label" (str), optional "variant": "error"
        kind = "task"         → "step" (int badge), "action" (bold), "description"
        kind = "gateway"      → "label" (str, rendered below the diamond)
        Optional on any node: "y_offset_in_lane" (px relative to lane centre;
                              negative = above centre, positive = below)

    Canvas dimensions auto-compute from the column count (derived from
    max node.col + 1) and the lane count. Embed at ~6.5 inches wide in
    the docx via `make_bpmn_diagram`.

    Returns the output_path so callers can chain `make_bpmn_diagram(doc,
    render_bpmn_image(data, path))`.
    """
    if not _PILLOW_AVAILABLE:
        raise ImportError(
            "Pillow is required for BPMN diagrams; install with "
            "`./v19/odoo/.venv/bin/pip install Pillow` or remove "
            "`bpmn_diagram` from the workflow's SPEC_DATA."
        )

    lanes = diagram_data.get("lanes") or []
    nodes = diagram_data.get("nodes") or []
    edges = diagram_data.get("edges") or []
    title = diagram_data.get("title", "")
    subtitle_lines = diagram_data.get("subtitle") or []
    show_legend = diagram_data.get("legend", True)

    if not lanes:
        raise ValueError("bpmn_diagram requires non-empty `lanes`.")
    if not nodes:
        raise ValueError("bpmn_diagram requires non-empty `nodes`.")

    # ----- Layout geometry (px @ _BPMN_DPI) ---------------------------
    # All dimensions are pixels at _BPMN_DPI, so on-page size ≈ px × DPI ratio
    # once embedded (see make_bpmn_diagram). Generous shapes + ~27-31 px labels
    # land at body-ish scale for a 4–5 column diagram and stay legible (~7 pt)
    # even for a wide 7-column one — the readability fix.
    col_count = max((n["col"] for n in nodes), default=0) + 1
    lane_count = len(lanes)

    PAD_L = 28                    # page pad left of the role sidebar
    PAD_R = 120                   # right pad (rightmost shape half + label room)
    SIDEBAR_W = 224               # solid role-band sidebar
    COL_PAD_L = 70                # sidebar edge → first column centre
    COL_W = 250                   # column centre-to-centre (comfortable task gap)
    TASK_W = 198
    EVENT_R = 33
    GATEWAY_S = 50

    # ----- Fonts (px) --------------------------------------------------
    title_font = _resolve_bpmn_font(52)
    subtitle_font = _resolve_bpmn_font(33)
    lane_name_font = _resolve_bpmn_font(37)
    lane_sub_font = _resolve_bpmn_font(25)
    action_font = _resolve_bpmn_font(30)
    description_font = _resolve_bpmn_font(25)
    gateway_font = _resolve_bpmn_font(28)
    edge_label_font = _resolve_bpmn_font(29)
    badge_font = _resolve_bpmn_font(26)
    legend_font = _resolve_bpmn_font(26)
    end_label_font = _resolve_bpmn_font(27)

    # Measuring draw (text metrics independent of the final canvas size, which
    # we can't know until task height is settled below).
    _mdraw = ImageDraw.Draw(Image.new("RGB", (8, 8)))

    def _wrap_text(text, font, max_width):
        words = text.split()
        if not words:
            return [""]
        lines, current = [], words[0]
        for word in words[1:]:
            trial = current + " " + word
            if _mdraw.textbbox((0, 0), trial, font=font)[2] <= max_width:
                current = trial
            else:
                lines.append(current)
                current = word
        lines.append(current)
        return lines

    def _text_h(font):
        bbox = _mdraw.textbbox((0, 0), "Ag", font=font)
        return bbox[3] - bbox[1]

    # ----- Task height sized to content (uniform across the diagram) --
    # Every task shares one height = tall enough for the wordiest action +
    # description, so descriptions never spill past the rounded-rectangle.
    A_LINE = _text_h(action_font) + 6
    D_LINE = _text_h(description_font) + 4
    max_a, max_d = 1, 0
    for n in nodes:
        if n.get("kind") == "task":
            max_a = max(max_a, len(_wrap_text(n.get("action", ""), action_font, TASK_W - 40)))
            if n.get("description"):
                max_d = max(max_d, len(_wrap_text(n["description"], description_font, TASK_W - 28)))
    # No upper cap: the box must always contain its wrapped action + description,
    # or text spills past the card (and, for a bottom-lane task, past the pool
    # boundary). Lane height tracks the task height so cards stay clear of the
    # lane edges and outside event/gateway labels still fit.
    TASK_H = max(140, 56 + max_a * A_LINE + (10 + max_d * D_LINE if max_d else 0) + 16)
    LANE_H = max(276, TASK_H + 120)

    # Title block height grows with the subtitle line count. In the docx flow
    # the title + subtitle are emitted as real text above the image (always
    # readable at true point size — see build_spec), so `draw_header=False`
    # collapses the in-image header to a slim top pad. Standalone renders keep
    # the baked header.
    sub_lines_wrapped = []
    if draw_header:
        sub_max_w = max(900, 6 * COL_W)
        for line in subtitle_lines:
            sub_lines_wrapped.extend(_wrap_text(line, subtitle_font, sub_max_w))
        title_h = (66 if title else 0) + len(sub_lines_wrapped) * 44
        if title or subtitle_lines:
            title_h += 26
        LANES_TOP = max(48, title_h)
    else:
        LANES_TOP = 34
    LEGEND_H = 132 if show_legend else 46

    pool_left = PAD_L
    sidebar_right = pool_left + SIDEBAR_W
    active_x0 = sidebar_right + COL_PAD_L          # first column centre
    W = active_x0 + (col_count - 1) * COL_W + TASK_W // 2 + PAD_R
    H = LANES_TOP + lane_count * LANE_H + LEGEND_H
    pool_right = W - 24

    col_x = [active_x0 + i * COL_W for i in range(col_count)]
    lane_top = [LANES_TOP + i * LANE_H for i in range(lane_count)]
    lane_bottom = [t + LANE_H for t in lane_top]
    lane_cy = [t + LANE_H // 2 for t in lane_top]

    R = _BPMN_RGB
    img = Image.new("RGB", (W, H), "white")
    draw = ImageDraw.Draw(img)

    def _draw_centered_lines(cx, top_y, lines, font, colour, line_gap=4,
                             min_x=None, max_x=None):
        line_h = _text_h(font) + line_gap
        y = top_y
        for line in lines:
            bbox = draw.textbbox((0, 0), line, font=font)
            tw = bbox[2] - bbox[0]
            x = cx - tw // 2
            # Keep labels inside the lane body — never spilling onto the dark
            # role sidebar or past the pool edge.
            if min_x is not None:
                x = max(x, min_x)
            if max_x is not None:
                x = min(x, max_x - tw)
            draw.text((x, y), line, fill=colour, font=font)
            y += line_h
        return y

    # ----- Title + subtitle (baked only for standalone renders) -------
    if draw_header:
        if title:
            draw.text((PAD_L + 4, 30), title, fill=R["ink"], font=title_font)
        sy = (74 if title else 28)
        for line in sub_lines_wrapped:
            draw.text((PAD_L + 6, sy), line, fill=R["secondary"], font=subtitle_font)
            sy += 44

    # ----- Swimlanes: alternating bodies, solid role sidebar, dividers -
    for i, lane in enumerate(lanes):
        top, bottom = lane_top[i], lane_bottom[i]
        # Lane body (alternating tint so the roles read as distinct bands).
        body_fill = "white" if i % 2 == 0 else R["lane_alt"]
        draw.rectangle([(sidebar_right, top), (pool_right, bottom)],
                       fill=body_fill, outline=None)
        # Solid role sidebar (dark plum band) with white role name + sub.
        draw.rectangle([(pool_left, top), (sidebar_right, bottom)],
                       fill=R["lane_label"], outline=None)
        name = lane.get("name", "")
        sub = lane.get("sub", "")
        name_lines = _wrap_text(name, lane_name_font, SIDEBAR_W - 36)
        sub_lines = _wrap_text(sub, lane_sub_font, SIDEBAR_W - 36) if sub else []
        nh = (_text_h(lane_name_font) + 6) * len(name_lines)
        sh = (_text_h(lane_sub_font) + 4) * len(sub_lines)
        block_h = nh + (10 + sh if sub_lines else 0)
        cx = pool_left + SIDEBAR_W // 2
        y = top + (LANE_H - block_h) // 2
        y = _draw_centered_lines(cx, y, name_lines, lane_name_font, "white", line_gap=6)
        if sub_lines:
            y += 10
            _draw_centered_lines(cx, y, sub_lines, lane_sub_font,
                                 (0xD8, 0xC8, 0xD2), line_gap=4)

    # Lane divider lines (between lanes) + sidebar seam + outer pool frame.
    for i in range(1, lane_count):
        draw.line([(pool_left, lane_top[i]), (pool_right, lane_top[i])],
                  fill=R["lane_divider"], width=2)
    draw.line([(sidebar_right, lane_top[0]), (sidebar_right, lane_bottom[-1])],
              fill=R["accent"], width=2)
    draw.rectangle([(pool_left, lane_top[0]), (pool_right, lane_bottom[-1])],
                   outline=R["accent"], width=3)

    # ----- Node position pre-pass + in-lane clamping ------------------
    # A node's centre is clamped so its full extent (shape + an outside-label
    # band where applicable) stays inside its own lane — fixes the error-end
    # event that used to float up past the swimlane boundary.
    LABEL_BAND = 58               # room reserved for an outside label (1–2 lines)
    EDGE = 16                     # min gap from the lane divider

    def _half_extent(kind):
        if kind == "task":
            return TASK_H // 2 + 6
        if kind == "gateway":
            return GATEWAY_S + LABEL_BAND
        return EVENT_R + LABEL_BAND       # start / end events

    node_by_id = {}
    for n in nodes:
        if n["lane"] >= lane_count:
            raise ValueError(f"node {n['id']!r}: lane {n['lane']} out of range")
        if n["col"] >= col_count:
            raise ValueError(f"node {n['id']!r}: col {n['col']} out of range")
        li = n["lane"]
        cx = col_x[n["col"]]
        raw_cy = lane_cy[li] + int(n.get("y_offset_in_lane", 0))
        half = _half_extent(n["kind"])
        lo = lane_top[li] + EDGE + half
        hi = lane_bottom[li] - EDGE - half
        if lo > hi:                        # lane too short for the band — centre it
            cy = lane_cy[li]
        else:
            cy = max(lo, min(hi, raw_cy))
        node_by_id[n["id"]] = dict(n, cx=cx, cy=cy, _li=li)

    # ----- Geometry validation: same-(lane, col) clearance check --------
    # Catches the bug class where two-or-more nodes share a cell without
    # enough vertical separation for their label bands. Author signal: an
    # author who stacked two end_events in the same column using only a
    # small `y_offset_in_lane` would see their labels overlap into
    # unreadable noise. This pass walks every (lane, col) cluster, sorts
    # by cy, and warns when adjacent nodes' label bands cross.
    # Required gap = (prev_half_extent + cur_half_extent + 16) where
    # half_extent is the node's full vertical reach including its label
    # band. Warn-only (not raise) so authors can override deliberately —
    # but every warning surfaces a real visual collision.
    from collections import defaultdict
    _cell_clusters = defaultdict(list)
    for n in nodes:
        _cell_clusters[(n["lane"], n["col"])].append(n)
    for (li, col), bucket in _cell_clusters.items():
        if len(bucket) < 2:
            continue
        bucket_sorted = sorted(bucket, key=lambda nn: node_by_id[nn["id"]]["cy"])
        for prev_n, cur_n in zip(bucket_sorted, bucket_sorted[1:]):
            prev_pos = node_by_id[prev_n["id"]]
            cur_pos = node_by_id[cur_n["id"]]
            prev_half = _half_extent(prev_n["kind"])
            cur_half = _half_extent(cur_n["kind"])
            gap = cur_pos["cy"] - prev_pos["cy"]
            required = prev_half + cur_half + 16  # +16: visual breathing room
            if gap < required:
                sys.stderr.write(
                    f"WARNING [bpmn geometry]: nodes {prev_n['id']!r} and "
                    f"{cur_n['id']!r} share (lane={li}, col={col}) with "
                    f"insufficient vertical clearance "
                    f"(gap={gap}px, required={required}px). "
                    f"Labels will likely overlap — split to adjacent columns "
                    f"(WF4 pattern: one node per col, no y_offset_in_lane), "
                    f"or increase y_offset_in_lane to at least "
                    f"{required - (gap - int(cur_n.get('y_offset_in_lane', 0)))}px.\n"
                )

    # ----- Shape drawing primitives -----------------------------------
    def _label_below_or_above(cx, cy, shape_half, lines, font, colour, li,
                              prefer="below"):
        """Place an event/gateway label clear of the shape. Default prefers the
        band below the shape; pass prefer="above" for gateways, whose branch
        edges (and their Yes/No labels) leave downward/sideways and would
        otherwise collide with a below-placed question label. Never crosses the
        lane boundary; clamped inside the lane body (off the role sidebar).

        `lo_x = sidebar_right + 22` (not +10) gives start/end labels at col=0
        a visible gap from the sidebar's right edge — at +10 a single-line
        label like "Reaches Pay step" would butt against the lane divider's
        2px-wide plum line and the eye reads it as clipped.
        """
        line_h = _text_h(font) + 4
        need = line_h * len(lines) + 8
        lo_x, hi_x = sidebar_right + 22, pool_right - 10
        below_top = cy + shape_half + 8
        above_top = cy - shape_half - need
        fits_below = below_top + need <= lane_bottom[li] - 4
        fits_above = above_top >= lane_top[li] + 4
        if prefer == "above":
            use_above = fits_above or not fits_below
        else:
            use_above = (not fits_below) and fits_above
        if use_above:
            _draw_centered_lines(cx, max(lane_top[li] + 4, above_top), lines,
                                 font, colour, min_x=lo_x, max_x=hi_x)
        else:
            _draw_centered_lines(cx, below_top, lines, font, colour,
                                 min_x=lo_x, max_x=hi_x)

    def draw_start(cx, cy, label, li):
        draw.ellipse([(cx - EVENT_R, cy - EVENT_R), (cx + EVENT_R, cy + EVENT_R)],
                     fill="white", outline=R["accent"], width=3)
        if label:
            _label_below_or_above(cx, cy, EVENT_R,
                                  _wrap_text(label, end_label_font, COL_W - 24),
                                  end_label_font, R["secondary"], li)

    def draw_end(cx, cy, variant, label, li):
        colour = R["warn"] if variant == "error" else R["success"]
        draw.ellipse([(cx - EVENT_R, cy - EVENT_R), (cx + EVENT_R, cy + EVENT_R)],
                     fill="white", outline=colour, width=6)
        if label:
            _label_below_or_above(cx, cy, EVENT_R,
                                  _wrap_text(label, end_label_font, COL_W - 24),
                                  end_label_font, colour, li)

    def draw_task(cx, cy, step, action, description):
        x0, y0 = cx - TASK_W // 2, cy - TASK_H // 2
        x1, y1 = cx + TASK_W // 2, cy + TASK_H // 2
        draw.rounded_rectangle([(x0, y0), (x1, y1)], radius=16,
                               fill="white", outline=R["accent"], width=3)
        # Accent top-stripe gives the task a recognisably-Odoo card feel.
        draw.rounded_rectangle([(x0, y0), (x1, y0 + 12)], radius=6,
                               fill=R["accent"], outline=R["accent"])
        if step is not None:
            badge_r = 19
            bx, by = x0 + badge_r + 12, y0 + badge_r + 16
            draw.ellipse([(bx - badge_r, by - badge_r), (bx + badge_r, by + badge_r)],
                         fill=R["accent"], outline="white", width=2)
            s = str(step)
            bbox = draw.textbbox((0, 0), s, font=badge_font)
            draw.text((bx - (bbox[2] - bbox[0]) // 2, by - (bbox[3] - bbox[1]) // 2 - bbox[1]),
                      s, fill="white", font=badge_font)
        action_lines = _wrap_text(action, action_font, TASK_W - 40)
        desc_lines = _wrap_text(description or "", description_font, TASK_W - 28) \
            if description else []
        # Top-align content below the badge/stripe so the step number never
        # overlaps the first action line.
        y = y0 + 56
        y = _draw_centered_lines(cx, y, action_lines, action_font, R["ink"], line_gap=6)
        if desc_lines:
            y += 8
            _draw_centered_lines(cx, y, desc_lines, description_font,
                                 R["secondary"], line_gap=4)

    def draw_gateway(cx, cy, label, li, prefer="above"):
        """Render the gateway diamond + X marker + label.

        `prefer` is computed by the caller from `gateway_prefer[id]` — see the
        precompute pass below the node-positioning loop. When the gateway has
        outgoing edges going to an UPPER lane (target lane index < source
        lane index), the branch labels render above the segments; an "above"
        gateway label would collide with them. The caller flips to "below" in
        that case so the question text sits cleanly under the diamond.
        """
        pts = [(cx, cy - GATEWAY_S), (cx + GATEWAY_S, cy),
               (cx, cy + GATEWAY_S), (cx - GATEWAY_S, cy)]
        draw.polygon(pts, fill="white")
        draw.line(pts + [pts[0]], fill=R["accent"], width=3)
        m = 17
        draw.line([(cx - m, cy - m), (cx + m, cy + m)], fill=R["accent"], width=4)
        draw.line([(cx - m, cy + m), (cx + m, cy - m)], fill=R["accent"], width=4)
        if label:
            _label_below_or_above(cx, cy, GATEWAY_S,
                                  _wrap_text(label.replace("\n", " "),
                                             gateway_font, COL_W - 12),
                                  gateway_font, R["ink"], li, prefer=prefer)

    # ----- Orthogonal edge routing (90° elbows) -----------------------
    def _hw_hh(pos):
        k = pos["kind"]
        if k == "task":
            return TASK_W // 2, TASK_H // 2
        if k == "gateway":
            return GATEWAY_S, GATEWAY_S
        return EVENT_R, EVENT_R

    def _anchor(pos, side):
        cx, cy = pos["cx"], pos["cy"]
        hw, hh = _hw_hh(pos)
        return {
            "left": (cx - hw, cy), "right": (cx + hw, cy),
            "top": (cx, cy - hh), "bottom": (cx, cy + hh),
        }[side]

    def _pick_label_seg(pts):
        """Pick the segment to hang an edge label on: the first segment long
        enough to clear an exit stub (>= 48 px), else the longest. Keeps a
        Yes/No label on a visible run near the branch instead of crammed onto a
        tiny stub against the source shape."""
        best, best_len = (pts[0], pts[1]), -1
        for i in range(len(pts) - 1):
            (x1, y1), (x2, y2) = pts[i], pts[i + 1]
            seg_len = abs(x2 - x1) + abs(y2 - y1)
            if seg_len >= 48:
                return (pts[i], pts[i + 1])
            if seg_len > best_len:
                best_len, best = seg_len, (pts[i], pts[i + 1])
        return best

    def _route(src, tgt):
        """Return (polyline_points, label_segment) routed with right angles.

        Cross-lane edges run their long vertical leg in a COLUMN GUTTER — the
        node-free band halfway between two column centres — so the riser can
        never pierce an intermediate node that shares the source or target
        column (e.g. a task stacked in another lane of the same column). Short
        stubs step the line off the source/target edge into the lane margin
        before the gutter run, so the connecting horizontals clear centred
        shapes too."""
        if src["_li"] == tgt["_li"]:
            if src["col"] < tgt["col"]:
                a, d = _anchor(src, "right"), _anchor(tgt, "left")
            elif src["col"] > tgt["col"]:
                a, d = _anchor(src, "left"), _anchor(tgt, "right")
            else:  # same column, stacked vertically
                if src["cy"] <= tgt["cy"]:
                    a, d = _anchor(src, "bottom"), _anchor(tgt, "top")
                else:
                    a, d = _anchor(src, "top"), _anchor(tgt, "bottom")
                return [a, d], (a, d)
            if abs(a[1] - d[1]) < 5:
                return [a, d], (a, d)
            mx = (a[0] + d[0]) // 2
            pts = [a, (mx, a[1]), (mx, d[1]), d]
            return pts, _pick_label_seg(pts)
        # Cross-lane: stub off the source edge, cross in a node-free gutter, stub
        # into the target. Gutter sits half a column to the source-side of the
        # target column, where no node centre ever lands.
        down = tgt["_li"] > src["_li"]
        a = _anchor(src, "bottom" if down else "top")
        d = _anchor(tgt, "top" if down else "bottom")
        if abs(a[0] - d[0]) < 5:                 # same column: straight drop
            return [a, d], (a, d)
        half = COL_W // 2
        gutter = d[0] - half if d[0] > a[0] else d[0] + half
        sli, tli = src["_li"], tgt["_li"]
        if down:
            sy = min(a[1] + 24, lane_bottom[sli] - 6)
            ty = max(d[1] - 24, lane_top[tli] + 6)
        else:
            sy = max(a[1] - 24, lane_top[sli] + 6)
            ty = min(d[1] + 24, lane_bottom[tli] - 6)
        pts = [a, (a[0], sy), (gutter, sy), (gutter, ty), (d[0], ty), d]
        return pts, _pick_label_seg(pts)

    def _draw_edge(pts):
        import math
        draw.line(pts, fill=R["accent"], width=3, joint="curve")
        (x1, y1), (x2, y2) = pts[-2], pts[-1]
        ang = math.atan2(y2 - y1, x2 - x1)
        head = 18
        draw.polygon([
            (x2, y2),
            (x2 - head * math.cos(ang - 0.42), y2 - head * math.sin(ang - 0.42)),
            (x2 - head * math.cos(ang + 0.42), y2 - head * math.sin(ang + 0.42)),
        ], fill=R["accent"])

    def _draw_edge_label(seg, label, label_colour):
        """Drawn in a final pass, ON TOP of shapes and their labels, so a
        Yes/No/branch label is never buried under a node or another label."""
        colour = R.get(label_colour or "", R["ink"])
        (sx1, sy1), (sx2, sy2) = seg
        mx, my = (sx1 + sx2) // 2, (sy1 + sy2) // 2
        bbox = draw.textbbox((0, 0), label, font=edge_label_font)
        tw, th = bbox[2] - bbox[0], bbox[3] - bbox[1]
        if abs(sx2 - sx1) >= abs(sy2 - sy1):     # horizontal segment
            ox, oy = -tw // 2, -th - 12
        else:                                    # vertical segment
            ox, oy = 12, -th // 2
        px, py = mx + ox, my + oy
        for dx in (-2, -1, 0, 1, 2):             # white halo for legibility
            for dy in (-2, -1, 0, 1, 2):
                if dx or dy:
                    draw.text((px + dx, py + dy), label, fill="white",
                              font=edge_label_font)
        draw.text((px, py), label, fill=colour, font=edge_label_font)

    # ----- Draw edge LINES first (node shapes sit on top of the line ends).
    # ----- Gateway label-placement preference (based on outgoing edges) ---
    # When a gateway's outgoing edges go to an UPPER lane (target lane index
    # smaller than the source's), the edges leave the gateway upward; their
    # Yes/No labels render above the segments and would collide with an
    # ABOVE-placed gateway label, producing the strike-through overlap visible
    # in early TEST007 WF2/WF3 renders. Flip to BELOW in that case so the
    # gateway's question sits cleanly under the diamond while the branch
    # labels keep their natural above-segment positions.
    gateway_prefer = {}
    for n in nodes:
        if n.get("kind") != "gateway":
            continue
        src_li = n["lane"]
        has_upward = False
        for e in edges:
            if e.get("from") != n["id"]:
                continue
            tgt = node_by_id.get(e["to"])
            if tgt and tgt["_li"] < src_li:
                has_upward = True
                break
        gateway_prefer[n["id"]] = "below" if has_upward else "above"

    # Edge LABELS are queued and drawn in a final pass on top of everything, so
    # a branch label can never be buried under a shape or a shape's own label.
    edge_label_queue = []
    for e in edges:
        src = node_by_id.get(e["from"])
        tgt = node_by_id.get(e["to"])
        if not src or not tgt:
            raise ValueError(
                f"edge references unknown node: from={e.get('from')!r} "
                f"to={e.get('to')!r}"
            )
        pts, seg = _route(src, tgt)
        _draw_edge(pts)
        if e.get("label"):
            edge_label_queue.append((seg, e["label"], e.get("label_colour")))

    # ----- Draw nodes --------------------------------------------------
    for n in nodes:
        pos = node_by_id[n["id"]]
        cx, cy, li = pos["cx"], pos["cy"], pos["_li"]
        kind = n["kind"]
        if kind == "start_event":
            draw_start(cx, cy, n.get("label"), li)
        elif kind == "end_event":
            draw_end(cx, cy, n.get("variant"), n.get("label"), li)
        elif kind == "task":
            draw_task(cx, cy, n.get("step"), n.get("action", ""),
                      n.get("description", ""))
        elif kind == "gateway":
            draw_gateway(cx, cy, n.get("label", ""), li,
                         prefer=gateway_prefer.get(n["id"], "above"))
        else:
            raise ValueError(f"unknown node kind: {kind!r}")

    # ----- Edge labels (top pass, above shapes + their labels) --------
    for seg, label, label_colour in edge_label_queue:
        _draw_edge_label(seg, label, label_colour)

    # ----- Legend strip ------------------------------------------------
    if show_legend:
        legend_y = H - LEGEND_H // 2 + 6
        x = pool_left + 8
        gap = 20
        r = 15

        def _legend_text(after_x, text):
            draw.text((after_x, legend_y - _text_h(legend_font) // 2 - 2),
                      text, fill=R["secondary"], font=legend_font)
            bbox = draw.textbbox((0, 0), text, font=legend_font)
            return after_x + (bbox[2] - bbox[0])

        draw.ellipse([(x, legend_y - r), (x + 2 * r, legend_y + r)],
                     fill="white", outline=R["accent"], width=3)
        x = _legend_text(x + 2 * r + 8, "Start") + gap
        draw.ellipse([(x, legend_y - r), (x + 2 * r, legend_y + r)],
                     fill="white", outline=R["success"], width=5)
        x = _legend_text(x + 2 * r + 8, "End — success") + gap
        draw.ellipse([(x, legend_y - r), (x + 2 * r, legend_y + r)],
                     fill="white", outline=R["warn"], width=5)
        x = _legend_text(x + 2 * r + 8, "End — error") + gap
        draw.rounded_rectangle([(x, legend_y - r), (x + 2 * r + 12, legend_y + r)],
                               radius=6, fill="white", outline=R["accent"], width=3)
        x = _legend_text(x + 2 * r + 20, "Task") + gap
        pts = [(x + r, legend_y - r), (x + 2 * r, legend_y),
               (x + r, legend_y + r), (x, legend_y)]
        draw.polygon(pts, fill="white")
        draw.line(pts + [pts[0]], fill=R["accent"], width=3)
        _legend_text(x + 2 * r + 8, "Decision")

    output_path.parent.mkdir(parents=True, exist_ok=True)
    img.save(output_path, format="PNG", dpi=(_BPMN_DPI, _BPMN_DPI))
    return output_path


def make_diagram_heading(doc: Document, title: str, subtitle_lines: list) -> None:
    """Emit a process-diagram title + subtitle as real docx text above the
    embedded image. Rendering them as text (rather than baking them into the
    PNG) keeps them at true point size — a wide swimlane PNG scaled to the
    content column would otherwise shrink the caption to a few points."""
    if title:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(title)
        run.font.size = Pt(12)
        run.font.color.rgb = INK_STRONG
        _apply_font(run, FONT_HEADING)
    for line in (subtitle_lines or []):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(1)
        p.paragraph_format.line_spacing = 1.1
        run = p.add_run(line)
        run.font.size = Pt(9.5)
        run.font.name = FONT_BODY
        run.font.color.rgb = INK_MUTED


def _diagram_width(image_path: Path) -> Inches:
    """On-page width for a rendered diagram: its native size (pixels ÷ the DPI
    it was saved at) capped to the content column. Embedding at native size is
    what keeps the labels at readable point sizes — the old fixed 6.5-inch embed
    of a high-pixel canvas crushed the type. Wide diagrams still clamp to the
    column width (and shrink), but no longer below their natural scale."""
    if _PILLOW_AVAILABLE:
        try:
            with Image.open(image_path) as im:
                dpi = (im.info.get("dpi") or (_BPMN_DPI, _BPMN_DPI))[0] or _BPMN_DPI
                native_in = im.width / float(dpi)
            return Inches(min(_BPMN_CONTENT_W_IN, native_in))
        except Exception:
            pass
    return Inches(_BPMN_CONTENT_W_IN)


def make_bpmn_diagram(doc: Document, image_path: Path,
                       caption: str = None) -> None:
    """Embed a BPMN PNG into the docx, centred, at its readable native width
    (capped to the content column), with an optional italic grey caption.

    Caller is responsible for ensuring `image_path` exists — typically
    by calling `render_bpmn_image()` first. In the standard `build_spec`
    flow the pair is always called together: if `render_bpmn_image`
    fails (e.g. Pillow missing in a bare Python venv), the build
    crashes loud at that step and never reaches here. The Pillow pre-
    flight is handled at the Diagram-choice checkpoint (see SKILL.md);
    by the time we get to `make_bpmn_diagram` the renderer has run and
    the PNG exists.
    """
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run()
    run.add_picture(str(image_path), width=_diagram_width(image_path))
    if caption:
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.paragraph_format.space_after = Pt(8)
        cap_run = cap.add_run(caption)
        cap_run.italic = True
        cap_run.font.name = FONT_BODY
        cap_run.font.size = Pt(9)
        cap_run.font.color.rgb = SECONDARY


# ---------------------------------------------------------------------------
# Phase marker
# ---------------------------------------------------------------------------

def add_phase_marker(doc: Document, text: str, phase: int = 2) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(f"▌ PHASE {phase} — {text}")
    run.bold = True
    run.font.name = FONT_BODY
    run.font.size = BODY_SIZE
    run.font.color.rgb = PHASE2


# ---------------------------------------------------------------------------
# Spec assembly — top-level
#
# No appendices. The Glossary, Acknowledgement, and Open Nits appendices
# were all removed per the conciseness pass. Tone rule #1 (business
# language first, technical identifiers parenthetical) keeps jargon out
# of body prose so a glossary is redundant; the docx is the spec, not a
# sign-off ceremony. See content_outline.md § Explicitly NOT present.
# ---------------------------------------------------------------------------

def build_spec(spec_data: dict, output_path: Path) -> None:
    """Build the docx end-to-end from a spec_data dict.

    spec_data shape (canonical — Stage 1 lint enforces the docx-visible
    parts; non-rendered keys like `dev_handoff` are passed through and
    consumed by `odoo-plan-development` directly):
        task_code, client_name, functionality_title, odoo_version,
        subtitle (optional; defaults to "Functional Specification"),
        metadata (dict: prepared_by, role, client, companion_to (optional),
                  scope, date, version, consultancy (optional — footer line;
                  derived from `role` if absent)),
        business_case (dict: problem, who, why, out_of_scope (optional —
                       renders as a fourth row in §2 when set)),
        apps_impacted (list of [app, objective, functional_changes]),
        new_apps (list of [app, objective, key_capabilities]),
        workflows (list of dicts with name, summary, subsections; plus
                   OPTIONAL `bpmn_diagram` OR `flow_strip` for the
                   per-workflow process visualisation; see SKILL.md
                   § BPMN diagram block for the bpmn_diagram schema).
                   Each workflow's `subsections` dict carries the 11
                   per-workflow blocks keyed by snake_case (see
                   `_render_workflow_subsections` for the canonical list).
        dev_handoff (optional dict — NOT rendered to the docx; consumed
                     by odoo-plan-development. See SKILL.md § Dev-Only
                     Hand-off to odoo-plan-development for the shape.
                     The builder ignores this key entirely; any unknown
                     top-level key is also ignored, so dev_handoff is a
                     convention, not a schema requirement).

    Five numbered top-level sections (plus Cover + Table of Contents
    front matter): Odoo Version, Business Case, Apps Impacted, Functional
    Layout, Per-Workflow Detail. No top-level Success Criteria (moved to
    per-workflow subsection 1), no Glossary appendix, no Context/Preface,
    no Type of Development, no Acknowledgement, no Open Nits — see
    reference/content_outline.md § Explicitly NOT present.
    """
    doc = new_document()

    # Running header (Odoo logo + task-code) + confidentiality footer.
    setup_page_furniture(doc, spec_data, output_path)

    # Cover page (unnumbered — front matter). The Odoo logo sits centred at the
    # top of the cover body; the running-header logo is suppressed on page 1.
    make_cover(
        doc,
        title=f"{spec_data['client_name']} — {spec_data['functionality_title']}",
        subtitle=spec_data.get("subtitle", "Functional Specification"),
        metadata=spec_data["metadata"],
        logo_path=_find_logo(output_path),
    )

    # Table of Contents (unnumbered front matter) — POPULATED on disk: titles +
    # page numbers, clickable, refreshed on open. Entries derive from the same
    # structure emitted below (collect_toc), so the TOC can't drift.
    make_toc(doc, collect_toc(spec_data))

    # 1. Odoo Version (just the version number, one line)
    make_heading(doc, "1. Odoo Version", level=1)
    make_paragraph(doc, str(spec_data["odoo_version"]))

    # 2. Business Case — concise 3-row table sourced from Q5 free-text brief.
    make_heading(doc, "2. Business Case", level=1)
    bc = spec_data.get("business_case") or {}
    business_case_rows = [
        ["Dimension", "Statement"],
        ["Problem", bc.get("problem", "")],
        ["Affected role", bc.get("who", "")],
        ["Why it matters", bc.get("why", "")],
    ]
    # Optional explicit out-of-scope row — rendered only when supplied, so it
    # never forces an empty cell on specs that don't bound scope this way.
    if bc.get("out_of_scope"):
        business_case_rows.append(["Out of scope", bc["out_of_scope"]])
    make_table(doc, business_case_rows)

    # 3. Apps Impacted / New Apps Proposed (both sub-tables 3 cols —
    # Objective + Functional changes / Key capabilities — to make the
    # customization required functionally explicit).
    make_heading(doc, "3. Apps Impacted / New Apps Proposed", level=1)
    if spec_data.get("apps_impacted"):
        make_heading(doc, "3.1 Impacted Standard Apps", level=2)
        rows = [["App", "Objective", "Functional changes"]] + spec_data["apps_impacted"]
        make_table(doc, rows)
    if spec_data.get("new_apps"):
        make_heading(doc, "3.2 New Custom Apps", level=2)
        rows = [["App", "Objective", "Key capabilities"]] + spec_data["new_apps"]
        make_table(doc, rows)

    # 4. Functional Layout (workflows table; NO block diagram at top level)
    make_heading(doc, "4. Functional Layout — Workflows", level=1)
    overview_rows = [["Workflow", "Summary (10–12 words)"]]
    for w in spec_data["workflows"]:
        overview_rows.append([w["name"], w["summary"]])
    make_table(doc, overview_rows)

    # 5. Per-Workflow Detail
    make_heading(doc, "5. Per-Workflow Detail", level=1)
    # BPMN PNGs (if any) land next to the builder so they travel with
    # the per-spec folder. output_path.parent is the spec folder root;
    # `_reference/` sits inside it.
    bpmn_dir = output_path.parent / "_reference"
    task_code = spec_data.get("task_code", "spec")
    for w_idx, w in enumerate(spec_data["workflows"], start=1):
        make_heading(doc, f"5.{w_idx} {w['name']}", level=2)
        # Process-flow visualisation, in order of richness:
        #   1. `bpmn_diagram` block → rendered BPMN swimlane PNG (preferred
        #      when the workflow has actors + a gateway or non-linear flow);
        #   2. `flow_strip` list of step labels → modern PIL-rendered
        #      PNG (rounded pills + step badges + arrows) when Pillow
        #      is available; falls back to the legacy table renderer
        #      when Pillow is missing (bare Python venv that declined
        #      the Pillow install gate at the Diagram-choice checkpoint).
        #   3. neither → no process-flow visualisation for this workflow.
        if w.get("bpmn_diagram"):
            bd = w["bpmn_diagram"]
            bpmn_path = bpmn_dir / f"{task_code}-wf{w_idx}-bpmn.png"
            # Title + subtitle as real docx text (readable at true point size),
            # then the swimlane image with its baked header suppressed.
            make_diagram_heading(doc, bd.get("title"), bd.get("subtitle"))
            render_bpmn_image(bd, bpmn_path, draw_header=False)
            make_bpmn_diagram(doc, bpmn_path, caption=bd.get("caption"))
        elif w.get("flow_strip") and len(w["flow_strip"]) >= 3:
            if _PILLOW_AVAILABLE:
                flow_path = bpmn_dir / f"{task_code}-wf{w_idx}-flow.png"
                render_flow_strip_image(w["flow_strip"], flow_path)
                make_flow_strip_diagram(doc, flow_path)
            else:
                # Legacy table fallback — used only when Pillow is
                # genuinely missing AND the user declined the install
                # gate at the Diagram-choice checkpoint.
                make_flow_strip(doc, w["flow_strip"])
        _render_workflow_subsections(doc, w["subsections"], w_idx)

    # No appendices. The docx ends at Section 5.

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))


# Canonical set of vague-N/A reasons. Single source of truth — the lint
# script (_lint_spec.py) imports this set so build-time and post-build
# checks stay aligned. Match semantics: exact match against the reason
# AFTER "N/A — ", lower-cased and stripped. The builder fails LOUD at
# write time when a vague reason is detected so the interview author
# notices immediately; the lint is defence-in-depth (catches builder-
# bypass paths, e.g. if a per-spec builder fork ever skips the check).
_VAGUE_NA_REASONS = frozenset({
    "not applicable",
    "n/a",
    "tbd",
    "to be determined",
    "to be defined",
    "not specified",
    "not specified yet",
    "unknown",
    "pending",
    "none",
    "no",
    "nothing",
})


def _render_workflow_subsections(doc: Document, subs: dict, w_idx: int) -> None:
    """Render the 11 per-workflow subsections in canonical order, numbered
    5.{w_idx}.{1..11} so they thread into the document's outline.

    Each subs[key] is either:
        {"table": [[...], ...]}    → table render (first row = header)
        {"content": "..."}         → prose paragraph
        {"na_reason": "..."}       → 'N/A — <reason>' (must be specific)
        {"sub_tables": {...}}      → multi-table render (used by Access &
                                     Permissions — see below)

    The `access` subsection specifically supports `sub_tables` with keys
    `new_groups` and `existing_groups`, each a regular table list. Either
    may be replaced by `na_reason` independently.
    """
    canonical_order = CANONICAL_SUBSECTIONS
    for sub_idx, (key, heading) in enumerate(canonical_order, start=1):
        sub = subs.get(key, {})
        make_heading(doc, f"5.{w_idx}.{sub_idx} {heading}", level=3)
        # Access & Permissions has two optional sub-tables (new groups +
        # existing groups mentioned) per the canonical shape.
        if key == "access" and sub.get("sub_tables"):
            _render_access_sub_tables(doc, sub["sub_tables"], w_idx, sub_idx)
            continue
        if "na_reason" in sub:
            reason = (sub["na_reason"] or "").strip()
            if reason.lower() in _VAGUE_NA_REASONS or not reason:
                raise ValueError(
                    f"Workflow {w_idx} subsection '{heading}' has a vague "
                    f"N/A reason: {reason!r}. Stage-1 lint rejects vague "
                    "reasons. Give a specific one (e.g. 'standard Odoo "
                    "credit-control flow unchanged for this workflow')."
                )
            make_paragraph(doc, f"N/A — {reason}")
        elif sub.get("table"):
            make_table(doc, sub["table"])
        elif sub.get("content"):
            make_paragraph(doc, sub["content"])
        else:
            # No content AND no na_reason — fail loud during build rather
            # than emit a placeholder that the lint will only catch later.
            raise ValueError(
                f"Workflow {w_idx} subsection '{heading}' has neither "
                "`table`, `content`, nor `na_reason`. Strict-N/A "
                "discipline: every subsection must be either filled or "
                "explicitly marked with a specific N/A reason."
            )


def _render_access_sub_tables(doc: Document, sub_tables: dict,
                              w_idx: int, sub_idx: int) -> None:
    """Render the Access & Permissions sub-tables — new groups + existing
    groups mentioned. Each may be a table list (header + rows) or carry
    an `na_reason` string when only one applies.
    """
    new_groups = sub_tables.get("new_groups")
    existing_groups = sub_tables.get("existing_groups")
    # 11.{w}.1 — New Groups
    make_heading(doc, f"5.{w_idx}.{sub_idx}.1 New Groups", level=4)
    if isinstance(new_groups, dict) and new_groups.get("na_reason"):
        reason = new_groups["na_reason"].strip()
        if reason.lower() in _VAGUE_NA_REASONS:
            raise ValueError(
                f"Workflow {w_idx} Access & Permissions / New Groups has a "
                f"vague N/A reason: {reason!r}."
            )
        make_paragraph(doc, f"N/A — {reason}")
    elif isinstance(new_groups, list) and new_groups:
        make_table(doc, new_groups)
    else:
        raise ValueError(
            f"Workflow {w_idx} Access & Permissions / New Groups must be "
            "either a table list or `{{na_reason: ...}}`."
        )
    # 11.{w}.2 — Existing Groups mentioned
    make_heading(doc, f"5.{w_idx}.{sub_idx}.2 Existing Groups mentioned",
                 level=4)
    if isinstance(existing_groups, dict) and existing_groups.get("na_reason"):
        reason = existing_groups["na_reason"].strip()
        if reason.lower() in _VAGUE_NA_REASONS:
            raise ValueError(
                f"Workflow {w_idx} Access & Permissions / Existing Groups "
                f"has a vague N/A reason: {reason!r}."
            )
        make_paragraph(doc, f"N/A — {reason}")
    elif isinstance(existing_groups, list) and existing_groups:
        make_table(doc, existing_groups)
    else:
        raise ValueError(
            f"Workflow {w_idx} Access & Permissions / Existing Groups must "
            "be either a table list or `{{na_reason: ...}}`."
        )


# ---------------------------------------------------------------------------
# CLI
# ---------------------------------------------------------------------------

def main() -> int:
    ap = argparse.ArgumentParser(description=__doc__.splitlines()[0])
    ap.add_argument("--output", required=True, help="Output .docx path")
    ap.add_argument("--spec-data", required=True, help="Path to a Python file defining SPEC_DATA dict")
    args = ap.parse_args()

    spec_data_path = Path(args.spec_data).resolve()
    if not spec_data_path.is_file():
        sys.stderr.write(f"spec-data file not found: {spec_data_path}\n")
        return 1
    namespace: dict = {}
    exec(spec_data_path.read_text(), namespace)
    spec_data = namespace.get("SPEC_DATA")
    if not isinstance(spec_data, dict):
        sys.stderr.write(f"{spec_data_path} must define a SPEC_DATA dict\n")
        return 1

    build_spec(spec_data, Path(args.output).resolve())
    print(f"Wrote {args.output}")
    return 0


if __name__ == "__main__":
    sys.exit(main())
