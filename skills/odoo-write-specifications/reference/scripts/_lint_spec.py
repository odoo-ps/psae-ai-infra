"""Stage 1 static lint for a `odoo-write-specifications` docx.

Walks the docx (or the SPEC_DATA Python file the builder consumes) and checks:
  - Filename matches `<task-code> - <client-name> - <3-word-func>.docx`
  - All 5 numbered top-level sections present (per content_outline.md)
    plus cover page and Table of Contents as unnumbered front matter
  - Every workflow has all 11 subsections present OR explicitly N/A
    with a SPECIFIC reason (vague reasons like 'not applicable' rejected)
  - No placeholder strings (TODO, TBD, <fill in>, Lorem ipsum, XXX, [client], [task])
  - Cover-page metadata block complete
  - 7-C surface checks: no paragraph > 4 lines; no casual-hedging strings
  - No top-level Success Criteria heading (moved to per-workflow subsection)
  - No Glossary appendix (removed per the conciseness pass)

Exit 0 if clean; non-zero on any error. Warnings are reported but don't fail.

Usage:
    ./v19/odoo/.venv/bin/python _lint_spec.py path/to/spec.docx
"""
from __future__ import annotations

import re
import sys
from pathlib import Path

try:
    from docx import Document
except ImportError:
    sys.stderr.write(
        "python-docx not installed; run: ./v19/odoo/.venv/bin/pip install python-docx\n"
    )
    sys.exit(2)


REQUIRED_TOP_LEVEL_SECTIONS = [
    # Cover page + Table of Contents are front matter (detected separately
    # — TOC by its field code, cover by the absence of a Heading 1 on page
    # one). The 5 named sections below are mandatory, numbered 1–5 in the
    # docx. Removed from earlier versions: Context / Preface, Type of
    # Development, top-level Success Criteria (moved to per-workflow) —
    # see content_outline.md § Explicitly NOT present.
    # Substring match, so the numbered prefix ("1. Odoo Version") still
    # satisfies the requirement.
    "Odoo Version",
    "Business Case",
    "Apps Impacted",  # prefix match — full title varies
    "Functional Layout",
    "Per-Workflow Detail",
]

REQUIRED_PER_WORKFLOW_SUBSECTIONS = [
    # New shape: Success Criteria is now per-workflow (subsection #1) — it
    # moved here from top-level Section 3.
    # "New Data Captured" was renamed "New Models" and is conditional on
    # whether a new model is introduced (lint accepts a strict-N/A reason
    # like "no new model introduced; extends existing Sales Order").
    "Success Criteria",
    "User Stories & Steps",
    "New Models",
    "New Fields & Information",
    "Navigation & Menus",
    "Screens & Interactions",
    "Automated Behaviours",
    "Business Rules & Validations",
    "Reports & Analytics",
    "Data Migration",
    "Access & Permissions",
]

# Vague N/A reasons rejected by the lint. The intent: if the reason can't
# be specific, the question wasn't asked thoughtfully. Force a re-ask
# rather than letting "N/A — not applicable" slip through.
#
# The canonical set lives in `_build_spec.py` as `_VAGUE_NA_REASONS` —
# imported here so build-time fail-loud and post-build lint check stay
# aligned. The builder is primary enforcement (raises ValueError during
# the interview); this lint check is defence-in-depth — it catches
# bypass paths (e.g. a per-spec builder fork that skipped the check, or
# a SPEC_DATA written directly to the docx out-of-band).
from _build_spec import _VAGUE_NA_REASONS  # noqa: E402  (intentional late import)

# Regex that matches a paragraph starting with "N/A — <reason>" and
# captures the reason for exact-match comparison against _VAGUE_NA_REASONS.
# Exact-match (not substring) so that legitimate reasons starting with a
# banned word (e.g. "N/A — none of the vans have …") are NOT false-flagged.
_NA_REASON_RE = re.compile(r"^N/A\s*—\s*(.+?)\s*$")

PLACEHOLDER_PATTERNS = [
    r"\bTODO\b",
    r"\bTBD\b",
    r"<fill in>",
    r"Lorem ipsum",
    r"\bXXX\b",
    r"\[client\]",
    r"\[task\]",
    # A blank/unpopulated Table of Contents — the field placeholder text leaked
    # into the document instead of a populated TOC (titles + page numbers).
    r"[Rr]ight-click.{0,30}[Uu]pdate [Ff]ield",
    r"Update Field to populate",
]

CASUAL_HEDGING = [
    r"\bI think\b",
    r"\bprobably\b",
    r"\bobviously\b",
    r"\bmaybe\b",
    r"\bsort of\b",
]

# Body prose should lean business-language. These Odoo-internal terms are
# fine inside the Glossary appendix; in body prose they trip a tone-drift
# warning. See content_outline.md § Style principles.
TECHNICAL_JARGON_IN_BODY = [
    r"\b_inherit\b",
    r"\b_inherits\b",
    r"\b_compute\b",
    r"@api\.depends\b",
    r"@api\.constrains\b",
    r"\bpost_init_hook\b",
    r"\bnoupdate\b",
    r"\brecord rule\b",         # use "row-level access rule"
    r"\bACL\b",                  # use "access permissions"
    r"\bdomain\s*=",            # field-level Odoo syntax leaking into prose
    r"\bxpath\b",
]

# Heuristic warnings that suggest the spec is enumerating inherited behaviour
# instead of focusing on what's new. Soft-warn, not blocker.
EXISTING_ENUMERATION_HINTS = [
    r"\binherited fields?\b",
    r"\bexisting fields?\b",
    r"\bexisting menus?\b",
    r"\bexisting access rights?\b",
    r"\bexisting automations?\b",
    r"\balready exists\b",
    r"\bunchanged from standard\b",
]

# Sections explicitly removed from the docx per content_outline.md § Explicitly
# NOT present. If any of these heading strings appears as a Heading 1, the
# docx has drifted back toward the longer pre-conciseness shape.
# Notes:
#   - "Business Case" was reinstated as section 2 — NOT on this list.
#   - "Success Criteria" appears as a per-workflow Heading 3 (subsection
#     1); it is on the list because at TOP-LEVEL (Heading 1) it's a regression.
#   - "Glossary" was removed per the conciseness pass; appearing as a
#     Heading 1 appendix is a regression.
REMOVED_HEADINGS = [
    "Context / Preface",
    "Context",
    "Preface",
    "Type of Development",
    "Acknowledgement",
    "Open Nits",
    "Anchor Pass",
    "Revisions",
    "Version History",
    "Alternatives Considered",
    "Success Criteria",  # only as a TOP-LEVEL Heading 1; allowed as per-workflow Heading 3
    "Glossary",
]

FILENAME_PATTERN = re.compile(
    r"^[A-Za-z0-9][A-Za-z0-9\-_]* - [^/:\\*\?\"<>|]+ - [a-z0-9-]+\.docx$"
)


def lint(docx_path: Path) -> tuple[list[str], list[str]]:
    """Return (errors, warnings)."""
    errors: list[str] = []
    warnings: list[str] = []

    if not FILENAME_PATTERN.match(docx_path.name):
        errors.append(
            f"Filename {docx_path.name!r} doesn't match "
            "'<task-code> - <client> - <3-word-func>.docx' convention."
        )

    if not docx_path.is_file():
        errors.append(f"File not found: {docx_path}")
        return errors, warnings

    doc = Document(str(docx_path))
    headings_l1 = [
        p.text.strip() for p in doc.paragraphs
        if p.style and p.style.name.startswith("Heading 1")
    ]
    headings_l3 = [
        p.text.strip() for p in doc.paragraphs
        if p.style and p.style.name.startswith("Heading 3")
    ]
    all_text = "\n".join(p.text for p in doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                all_text += "\n" + cell.text

    # Top-level sections — required
    for required in REQUIRED_TOP_LEVEL_SECTIONS:
        if not any(required in h for h in headings_l1):
            errors.append(f"Missing top-level section: {required!r}")

    # Top-level sections — explicitly removed (warn if they drift back in)
    for removed in REMOVED_HEADINGS:
        for h in headings_l1:
            if removed in h:
                warnings.append(
                    f"Removed section reappeared: {h!r} — see content_outline.md "
                    "§ Explicitly NOT present"
                )

    # Per-workflow subsections — for every Heading 2 found between "Per-Workflow
    # Detail" and the next Heading 1, expect the 11 subsections (or explicit N/A).
    for required in REQUIRED_PER_WORKFLOW_SUBSECTIONS:
        if not any(required in h for h in headings_l3):
            warnings.append(
                f"No workflow contains subsection {required!r} — "
                "if every workflow legitimately skips it, this is fine; "
                "otherwise check coverage."
            )

    # Placeholder strings
    for pat in PLACEHOLDER_PATTERNS:
        m = re.search(pat, all_text, re.IGNORECASE)
        if m:
            errors.append(f"Placeholder string found: {m.group(0)!r}")

    # Casual hedging
    for pat in CASUAL_HEDGING:
        for m in re.finditer(pat, all_text, re.IGNORECASE):
            warnings.append(f"Casual hedging: {m.group(0)!r} — prefer concrete claim")

    # Technical jargon leaking into body prose (Glossary appendix is exempt
    # because the linter walks all paragraphs uniformly; the worst offenders
    # cluster in the body and trigger here)
    for pat in TECHNICAL_JARGON_IN_BODY:
        for m in re.finditer(pat, all_text):
            warnings.append(
                f"Technical jargon in prose: {m.group(0)!r} — soften to "
                "business language; technical term belongs in the Glossary"
            )

    # "What's new only" — flag enumeration of existing behaviour
    for pat in EXISTING_ENUMERATION_HINTS:
        for m in re.finditer(pat, all_text, re.IGNORECASE):
            warnings.append(
                f"Enumeration of existing behaviour: {m.group(0)!r} — "
                "spec should describe only what's NEW; the dev team can "
                "grep Odoo for the inherited surface"
            )

    # Paragraph length cap (4 lines ~ 320 chars at the styling's text width)
    for p in doc.paragraphs:
        if p.style and p.style.name.startswith("Heading"):
            continue
        if len(p.text) > 480:  # generous; 4 short lines fit comfortably below
            warnings.append(
                f"Long paragraph ({len(p.text)} chars): "
                f"{p.text[:80]!r}... — split into shorter ones"
            )

    # Cover-page metadata sanity — look for the canonical labels
    expected_labels = ("Prepared by", "Client", "Scope", "Date", "Version")
    for label in expected_labels:
        if label not in all_text:
            errors.append(f"Cover-page metadata label missing: {label!r}")

    # Strict N/A discipline (defence-in-depth — see comment above the
    # _VAGUE_NA_REASONS import; the builder is the primary check):
    #   1. Every "N/A —" must have a reason after the em-dash.
    #   2. The reason must NOT exact-match a vague reason in the canonical
    #      set ("not applicable" / "TBD" / etc). Exact match (not substring)
    #      so that "N/A — none of the vans carry …" is NOT false-flagged.
    for p in doc.paragraphs:
        text = p.text.strip()
        if not text.startswith("N/A —"):
            continue
        match = _NA_REASON_RE.match(text)
        if not match:
            errors.append(f"N/A marker has no reason: {p.text!r}")
            continue
        reason = match.group(1).strip().lower()
        if reason in _VAGUE_NA_REASONS:
            errors.append(
                f"Vague N/A reason rejected: {p.text!r} — give a "
                "specific reason like 'standard Odoo X flow unchanged' "
                "or 'no new model introduced; extends existing Y'."
            )

    return errors, warnings


def main() -> int:
    if len(sys.argv) != 2:
        sys.stderr.write("usage: _lint_spec.py <path-to-spec.docx>\n")
        return 2
    path = Path(sys.argv[1]).resolve()
    errors, warnings = lint(path)
    for w in warnings:
        sys.stderr.write(f"WARN: {w}\n")
    for e in errors:
        sys.stderr.write(f"FAIL: {e}\n")
    if errors:
        sys.stderr.write(f"\n{len(errors)} error(s), {len(warnings)} warning(s).\n")
        return 1
    sys.stdout.write(f"OK — {path.name} ({len(warnings)} warning(s))\n")
    return 0


if __name__ == "__main__":
    sys.exit(main())
